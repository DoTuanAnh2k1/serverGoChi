"""Generate cli-config reference document (.docx).

Run: python3 docs/build_cli_config_doc.py
Output: docs/cli-config-reference.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ─────────────────────────────────────────────────────────────

def set_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    return p


def add_h(doc, text, level):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            t.rows[ri].cells[ci].text = v
    return t


# ── document ─────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_style(doc)

    title = doc.add_heading("cli-config — Hướng dẫn đầy đủ các lệnh", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("Tham chiếu toàn bộ verb / entity của REPL cli-config trong SSH bastion")
    srun.italic = True
    srun.font.size = Pt(11)

    # 1. Tổng quan
    add_h(doc, "1. Tổng quan", 1)
    add_para(doc,
             "cli-config là REPL chạy bên trong mode `cli-config` của SSH bastion (cli-gate). "
             "Mode này CHỈ hiển thị cho user role admin/superadmin — Normal user khi SSH vào "
             "chỉ thấy ne-config + ne-command. cli-config gọi tới mgt-svc qua HTTP+JWT để "
             "quản trị toàn bộ user, NE, group, RBAC, password policy, mgt permission. "
             "Tab để auto-complete; '--help' / '-h' append vào bất kỳ lệnh để in help; "
             "'exit' / 'quit' để quay về menu mode.")

    add_h(doc, "Cú pháp chung", 2)
    add_code(doc,
             "cli-config> <verb> <entity> [<target>] [<field> <value> ...]\n\n"
             "verb    : show | set | update | delete | purge | map | unmap |\n"
             "          allow | deny | revoke | help | exit\n"
             "entity  : user | ne | group | ne-profile | command-def | command-group\n"
             "target  : tên (name) hoặc id (số) của record\n"
             "field   : tên trường (có alias — xem bảng trường)\n"
             "value   : giá trị; nếu chứa khoảng trắng phải bỏ trong \"...\"")

    add_h(doc, "Quy tắc chung", 2)
    add_bullet(doc, "Field–value đi thành cặp, phân tách bằng khoảng trắng.")
    add_bullet(doc, "Giá trị có khoảng trắng phải đặt trong dấu ngoặc kép: full_name \"Alice Wonder\".")
    add_bullet(doc, "Mỗi entity có danh sách field BẮT BUỘC (cho `set`) và field TUỲ CHỌN.")
    add_bullet(doc, "account_type chỉ chấp nhận 1 (Admin) hoặc 2 (Normal). SuperAdmin không tạo được qua CLI.")
    add_bullet(doc, "conf_mode của NE chỉ chấp nhận: SSH, TELNET, NETCONF, RESTCONF.")
    add_bullet(doc, "Tab để auto-complete verb, entity, tên field, giá trị enum.")
    add_bullet(doc, "Append '--help' (hoặc '-h') vào bất kỳ vị trí trên dòng lệnh để in help context-aware "
                    "(`set user --help` in help cho `set user`; `show ne --help` cho `show ne`).")
    add_bullet(doc, "Lệnh thành công in `OK: <mô tả>`. Lệnh lỗi in `error: <chi tiết>` và giữ prompt.")
    add_bullet(doc, "Mọi lệnh `delete <entity> <target>` và `purge user <name>` đều bắt y/N confirm.")

    # 2. Bảng entities + fields
    add_h(doc, "2. Bảng trường cho từng entity (cho `set` / `update`)", 1)

    add_h(doc, "2.1. user (tbl_account)", 2)
    add_para(doc, "Field bắt buộc khi set:", bold=True)
    add_bullet(doc, "name (alias: account_name, username) — tên đăng nhập, duy nhất.")
    add_bullet(doc, "password — mgt-svc sẽ hash bcrypt trước khi lưu.")
    add_para(doc, "Field tuỳ chọn:", bold=True)
    add_table(doc,
              ["Alias", "Canonical", "Kiểu", "Ghi chú"],
              [
                  ["email", "email", "string", "Unique giữa các user active"],
                  ["full_name", "full_name", "string", "Bắt buộc nếu type=1 (Admin)"],
                  ["phone / phone_number", "phone_number", "string", "Bắt buộc nếu type=1; format E.164"],
                  ["address", "address", "string", "—"],
                  ["description", "description", "string", "—"],
                  ["type / account_type", "account_type", "int", "1=Admin, 2=Normal (0=SuperAdmin bị reject)"],
              ])

    add_h(doc, "2.2. ne (cli_ne)", 2)
    add_para(doc, "Field bắt buộc khi set:", bold=True)
    add_bullet(doc, "ne_name (alias: name)")
    add_bullet(doc, "namespace — (ne_name, namespace) duy nhất")
    add_bullet(doc, "conf_master_ip (alias: ip)")
    add_bullet(doc, "conf_port_master_tcp (alias: port)")
    add_bullet(doc, "command_url")
    add_para(doc, "Field tuỳ chọn (gồm `ne_profile` mới):", bold=True)
    add_table(doc,
              ["Alias", "Canonical", "Kiểu", "Ghi chú"],
              [
                  ["site / site_name", "site_name", "string", "Khu vực"],
                  ["system_type", "system_type", "string", "Tự do"],
                  ["description", "description", "string", "—"],
                  ["mode / conf_mode", "conf_mode", "enum", "SSH | TELNET | NETCONF | RESTCONF"],
                  ["conf_slave_ip", "conf_slave_ip", "string", "HA"],
                  ["conf_port_master_ssh", "conf_port_master_ssh", "int", "Cổng SSH master"],
                  ["conf_port_slave_ssh", "conf_port_slave_ssh", "int", "Cổng SSH slave"],
                  ["conf_port_slave_tcp", "conf_port_slave_tcp", "int", "Cổng TCP slave"],
                  ["conf_username", "conf_username", "string", "User SSH vào NE"],
                  ["conf_password", "conf_password", "string", "Pass SSH vào NE"],
                  ["profile / ne_profile", "ne_profile_id", "int|name", "Tên hoặc id của NE Profile"],
              ])

    add_h(doc, "2.3. group (cli_group)", 2)
    add_bullet(doc, "Bắt buộc: name. Tuỳ chọn: description.")
    add_bullet(doc, "`update group <name|id> name <new>` để rename. Permission gắn vào group bằng allow/deny (xem mục 3.6).")

    add_h(doc, "2.4. ne-profile (cli_ne_profile)", 2)
    add_bullet(doc, "Bắt buộc: name. Tuỳ chọn: description.")
    add_bullet(doc, "Phân loại NE theo tập lệnh (SMF, AMF, UPF, generic-router, …). NE gắn profile bằng `update ne <name> ne_profile <profile>`.")

    add_h(doc, "2.5. command-def (cli_command_def)", 2)
    add_para(doc, "Bắt buộc: service, pattern, category. Tuỳ chọn: ne_profile (default *), risk_level (0/1/2), description.", bold=True)
    add_table(doc,
              ["Alias", "Canonical", "Kiểu / Enum", "Ghi chú"],
              [
                  ["service", "service", "ne-command|ne-config|*", "Service nào áp dụng"],
                  ["profile / ne_profile", "ne_profile", "string", "Profile mà command thuộc về (* = mọi NE)"],
                  ["pattern", "pattern", "string", "Pattern lệnh; `*` cuối = wildcard suffix"],
                  ["category", "category", "monitoring|configuration|admin|debug", "—"],
                  ["risk / risk_level", "risk_level", "int 0/1/2", "0=safe, 2=dangerous"],
                  ["description", "description", "string", "—"],
              ])

    add_h(doc, "2.6. command-group (cli_command_group)", 2)
    add_bullet(doc, "Bắt buộc: name. Tuỳ chọn: service (default *), ne_profile (default *), description.")
    add_bullet(doc, "Quản lý member bằng `map command-group <cg> command <cmd_def_id>` / `unmap`.")

    # 3. Verbs
    add_h(doc, "3. Các verb chi tiết", 1)

    # 3.1 show
    add_h(doc, "3.1. show — liệt kê / xem chi tiết / lọc theo trường", 2)
    add_code(doc,
             "show <entity>                                # liệt kê tất cả\n"
             "show <entity> <name|id>                      # legacy: tìm theo name hoặc id\n"
             "show <entity> <field> <value>                # lọc theo field")
    add_para(doc, "Filter fields hỗ trợ:", bold=True)
    add_table(doc,
              ["Entity", "Filter field (alias)", "Hành vi"],
              [
                  ["user",
                   "name (username, account_name) | id | email | role (type)",
                   "name/id/email match chính xác → in detail (table khi trùng); role nhận label SuperAdmin|Admin|Normal hoặc 0|1|2 → in bảng."],
                  ["ne",
                   "name (ne_name) | id | site (site_name) | namespace",
                   "name in bảng nếu trùng qua nhiều namespace; site/namespace luôn in bảng."],
                  ["group", "name | id", "In detail (kèm users + ne_ids)."],
                  ["ne-profile", "name | id", "In detail."],
                  ["command-def", "service | ne_profile | category", "Filter, in bảng."],
                  ["command-group", "service | ne_profile | name", "name in detail (kèm member); còn lại in bảng."],
              ])
    add_para(doc, "Tab ở vị trí field gợi ý alias; Tab ở vị trí value của trường enum (role, service, category) gợi ý giá trị.")

    add_h(doc, "Ví dụ", 3)
    add_code(doc,
             "cli-config> show user role Admin\n"
             "ID  NAME    ROLE   ENABLED  EMAIL              FULL NAME\n"
             "2   alice   Admin  true     alice@example.com  Alice Wonder\n"
             "(1 user)")

    # 3.2 set
    add_h(doc, "3.2. set — tạo mới (hoặc re-enable nếu user đang disable)", 2)
    add_code(doc, "set <entity> <field> <value> [<field> <value> ...]")
    add_para(doc, "Tất cả field bắt buộc của entity phải có mặt — thiếu sẽ trả lỗi và huỷ lệnh.")
    add_para(doc, "Re-enable user disabled (đặc biệt cho `set user`):", bold=True)
    add_bullet(doc, "Nếu account_name đã tồn tại nhưng `is_enable=false`, server merge các field non-empty từ request vào record cũ + bật lại `is_enable=true`.")
    add_bullet(doc, "Field không truyền giữ nguyên giá trị cũ. Password luôn refresh.")
    add_bullet(doc, "Email của user disabled coi như free → user mới có thể tái sử dụng.")

    add_h(doc, "Ví dụ", 3)
    add_code(doc,
             "cli-config> set user name alice password Secret123! email alice@example.com \\\n"
             "            full_name \"Alice Wonder\" phone 0900000000 type 1\n"
             "OK: user created\n"
             "\n"
             "cli-config> set ne ne_name HTSMF01 namespace default ip 10.0.0.10 \\\n"
             "            port 830 command_url http://10.0.0.10/restconf mode NETCONF \\\n"
             "            site_name HN profile SMF\n"
             "OK: NE created\n"
             "\n"
             "cli-config> set ne-profile name SMF description \"Session Management Function\"\n"
             "OK: ne-profile created\n"
             "\n"
             "cli-config> set command-def service ne-command ne_profile SMF \\\n"
             "            pattern \"get subscriber\" category monitoring\n"
             "OK: command-def created")

    # 3.3 update
    add_h(doc, "3.3. update — sửa field", 2)
    add_code(doc, "update <entity> <name|id> <field> <value> [<field> <value> ...]")
    add_para(doc, "Chỉ truyền field cần đổi. Đặc biệt: `update ne <name> ne_profile <profile>` để gán/đổi profile cho NE.")
    add_code(doc,
             "cli-config> update user alice email new@example.com phone 0911111111\n"
             "OK: user updated\n"
             "cli-config> update ne HTSMF01 ne_profile SMF\n"
             "OK: NE updated")

    # 3.4 delete + purge
    add_h(doc, "3.4. delete (soft) vs purge (hard)", 2)
    add_para(doc, "Hai cấp xoá khác nhau — chọn đúng tuỳ trường hợp:", bold=True)
    add_table(doc,
              ["Verb", "Phạm vi", "Khôi phục", "Trường hợp dùng"],
              [
                  ["delete user <name>", "Set is_enable=false (soft delete)", "Có — gọi `set user name <same> password <…>` để re-enable",
                   "Tạm thời ngắt access, cần giữ history & có thể bật lại."],
                  ["purge user <name>", "Xoá hẳn row tbl_account + mọi mapping + password history",
                   "KHÔNG khôi phục được", "Cleanup test account, credential bị compromise, GDPR-style erasure. cli_operation_history vẫn được giữ."],
                  ["delete ne <name|id>", "Xoá NE (cascade user-ne, group-ne mappings)", "Không", "Bốc NE khỏi hệ thống."],
                  ["delete group <name|id>", "Xoá group + user-group, group-ne mappings", "Không", "—"],
                  ["delete ne-profile <name|id>", "Xoá profile (NE giữ profile_id stale)", "Không", "Sửa profile_id của NE liên quan trước."],
                  ["delete command-def <id>", "Xoá command def (cascade khỏi command-group)", "Không", "—"],
                  ["delete command-group <name|id>", "Xoá command group + member mapping", "Không", "Permission rule trỏ tới group này sẽ rỗng → coi như implicit deny."],
              ])
    add_para(doc, "Cả delete và purge đều bắt y/N confirm. Phần purge UI hiện thông báo `Delete PURGE user \"<name>\"? [y/N]`.")
    add_bullet(doc, "SuperAdmin: cả `delete user` và `purge user` đều bị refuse (403).")

    add_h(doc, "Ví dụ", 3)
    add_code(doc,
             "cli-config> delete user alice\n"
             "Delete user \"alice\"? [y/N]: y\n"
             "OK: user deleted\n"
             "\n"
             "cli-config> purge user alice\n"
             "Delete PURGE user \"alice\"? [y/N]: y\n"
             "OK: user purged (row + mappings + password history removed)")

    # 3.5 map / unmap
    add_h(doc, "3.5. map / unmap — quan hệ", 2)
    add_code(doc,
             "map user <username> ne <ne_name|id>             # gán NE trực tiếp cho user\n"
             "map user <username> group <group_name|id>       # cho user vào group\n"
             "map group <group_name|id> ne <ne_name|id>       # gán NE cho group\n"
             "map command-group <cg> command <cmd_def_id>     # add member vào bundle\n"
             "unmap ...                                       # cùng shape, ngược chiều")
    add_para(doc, "User có quyền lên NE qua 2 đường: (a) map user trực tiếp, (b) user thuộc group, group được map với NE. Unmap chỉ gỡ đúng quan hệ chỉ định.")

    # 3.6 allow / deny / revoke
    add_h(doc, "3.6. allow / deny / revoke — phân quyền lệnh trên NE (RBAC)", 2)
    add_code(doc,
             "allow  <group> <grant_type> <grant_value> [ne_scope <scope>] [service <svc>]\n"
             "deny   <group> <grant_type> <grant_value> [ne_scope <scope>] [service <svc>]\n"
             "revoke <group> <perm_id>")
    add_para(doc, "grant_type ∈ {command-group, category, pattern}. ne_scope ∈ {\"*\", \"profile:<name>\", \"ne:<ne_name>\"}.")
    add_para(doc, "Evaluation logic (AWS-IAM × Vault scope-specificity):", bold=True)
    add_bullet(doc, "Tại scope cụ thể nhất có rule match: explicit deny > explicit allow > implicit deny.")
    add_bullet(doc, "Scope priority: ne:X > profile:Y > * — rule cụ thể hơn override rule rộng hơn.")
    add_bullet(doc, "Pattern matching: `*` cuối = wildcard suffix; không có `*` = exact match hoặc prefix theo whitespace.")

    add_h(doc, "Ví dụ", 3)
    add_code(doc,
             "# Cho team-smf-l1 chỉ chạy được subscriber ops trên SMF; cấm session ops\n"
             "cli-config> allow team-smf-l1 command-group smf-subscriber-ops ne_scope profile:SMF\n"
             "cli-config> deny  team-smf-l1 command-group smf-session-ops    ne_scope profile:SMF\n"
             "\n"
             "# Cấm DELETE bất kỳ trên NE production cụ thể\n"
             "cli-config> deny noc-admin pattern \"delete *\" ne_scope ne:SMF-PROD-01\n"
             "\n"
             "# Xoá 1 rule theo id\n"
             "cli-config> revoke team-smf-l1 42\n"
             "Delete permission \"42\"? [y/N]: y\n"
             "OK: permission revoked")

    # 3.7 help
    add_h(doc, "3.7. help / --help", 2)
    add_code(doc,
             "help                          # tổng quan\n"
             "help <verb>                   # chi tiết 1 verb\n"
             "help <verb> <entity>          # chi tiết verb+entity (vd 'help set user')\n"
             "<bất kỳ lệnh> --help          # context-aware\n"
             "<bất kỳ lệnh> -h              # alias ngắn")

    # 3.8 exit
    add_h(doc, "3.8. exit / quit", 2)
    add_para(doc, "Thoát mode cli-config, quay lại menu `mode>`. Không nhận tham số.")

    # 4. Authorization model summary
    add_h(doc, "4. Mô hình authorization (đọc nhanh)", 1)
    add_para(doc, "User có quyền chạy lệnh Z trên NE Y trong service S khi:")
    add_bullet(doc, "(1) Y nằm trong tập NE access của user (union direct user-ne ∪ group-ne của mọi group user thuộc).")
    add_bullet(doc, "(2) Có rule allow trên một group user thuộc match (S, ne_scope phủ Y, grant_value match Z).")
    add_bullet(doc, "(3) KHÔNG có rule deny tại scope cụ thể nhất có match.")
    add_para(doc, "Downstream service (cli-netconf, cli-ne-command) gọi `GET /aa/authorize/rbac/effective` "
                  "để cache rule per session, hoặc `POST /aa/authorize/rbac/check-command` realtime cho lệnh risky.")

    # 5. Kịch bản end-to-end
    add_h(doc, "5. Kịch bản admin end-to-end", 1)
    add_code(doc,
             "$ ssh anhdt195@gate.example.com -p 2223\n"
             "Welcome anhdt195 — management CLI.\n"
             "Available modes: cli-config, ne-config, ne-command\n"
             "\n"
             "mode> cli-config\n"
             "== cli-config mode ==\n"
             "\n"
             "# 1) Tạo profile + NE và gán\n"
             "cli-config> set ne-profile name SMF\n"
             "cli-config> set ne ne_name HTSMF01 namespace default ip 10.0.0.10 port 830 \\\n"
             "            command_url http://10.0.0.10/restconf profile SMF\n"
             "\n"
             "# 2) Khai báo lệnh + nhóm lệnh\n"
             "cli-config> set command-def service ne-command ne_profile SMF \\\n"
             "            pattern \"get subscriber\" category monitoring\n"
             "cli-config> set command-def service ne-command ne_profile SMF \\\n"
             "            pattern \"delete session *\" category admin risk 2\n"
             "cli-config> set command-group name smf-subscriber-ops ne_profile SMF\n"
             "cli-config> map command-group smf-subscriber-ops command 1\n"
             "\n"
             "# 3) Tạo password policy + group + user\n"
             "# (password policy + assign group: dùng API /aa/password-policy + /aa/group/{id}/password-policy)\n"
             "cli-config> set group name team-smf description \"SMF operators\"\n"
             "cli-config> set user name alice password Secret123! type 1 \\\n"
             "            full_name \"Alice Wonder\" phone 0900000000\n"
             "cli-config> map user alice group team-smf\n"
             "cli-config> map group team-smf ne HTSMF01\n"
             "\n"
             "# 4) Phân quyền lệnh\n"
             "cli-config> allow team-smf command-group smf-subscriber-ops ne_scope profile:SMF\n"
             "cli-config> deny  team-smf pattern \"delete *\" ne_scope ne:HTSMF01\n"
             "\n"
             "# 5) Soft-delete vs purge\n"
             "cli-config> delete user alice            # tạm khoá\n"
             "cli-config> purge user alice             # xoá hẳn (sau khi đã chắc)\n"
             "\n"
             "cli-config> exit\n"
             "mode> exit")

    # 6. SSH gate role-based mode menu
    add_h(doc, "6. Menu mode theo role", 1)
    add_table(doc,
              ["Role (account_type)", "Mode hiển thị trong menu", "Ghi chú"],
              [
                  ["SuperAdmin (0) / Admin (1)", "cli-config, ne-config, ne-command", "Truy cập đầy đủ"],
                  ["Normal (2)", "ne-config, ne-command", "cli-config bị ẩn khỏi menu — chỉ admin mới thao tác cli-config được"],
              ])
    add_para(doc, "Sai password 3 lần liên tiếp (theo policy) sẽ bị lockout — server trả 403 `{locked_until, retry_in_seconds}`. Login đúng reset counter.")

    # 7. Errors
    add_h(doc, "7. Lỗi thường gặp", 1)
    add_table(doc,
              ["Thông báo", "Nguyên nhân", "Cách xử lý"],
              [
                  ["unknown command \"x\"", "Verb không hợp lệ", "Tab hoặc 'help' để xem verb hỗ trợ"],
                  ["unknown field \"x\" for user (valid: ...)", "Alias trường không hợp lệ", "Tab hoặc xem mục 2"],
                  ["missing required field \"X\" for set ne (required: ...)", "Thiếu field bắt buộc", "Bổ sung đủ field"],
                  ["field \"conf_mode\" must be one of [SSH TELNET NETCONF RESTCONF], got \"...\"", "Giá trị enum sai", "Chọn 1 trong các giá trị cho phép"],
                  ["field \"port\" must be an integer, got \"abc\"", "Field số nhập chuỗi", "Truyền số nguyên"],
                  ["no user with name or id \"x\"", "Target không tồn tại", "show <entity> để kiểm tra"],
                  ["unterminated quoted string", "Quote chưa đóng", "Đóng \" hoặc bỏ quote"],
                  ["refusing to purge SuperAdmin account", "Cố purge user account_type=0", "SuperAdmin chỉ xoá được trực tiếp trong DB"],
                  ["password too short / does not meet complexity", "Password vi phạm effective policy", "Đổi pass theo yêu cầu policy"],
                  ["account locked, retry after N seconds", "Sai pass nhiều lần", "Chờ hết lockout_minutes hoặc admin reset"],
              ])

    # 8. Tips
    add_h(doc, "8. Mẹo sử dụng", 1)
    add_bullet(doc, "Tab cycling: nhấn Tab liên tục để quay vòng candidate; nếu line dài vượt chiều ngang terminal, hint preview tự ẩn (cycling vẫn hoạt động).")
    add_bullet(doc, "Hint hiển thị ngay dòng dưới prompt khi có ≥2 candidate.")
    add_bullet(doc, "Alias: dùng 'name', 'ip', 'port', 'mode', 'type', 'profile' để gõ nhanh.")
    add_bullet(doc, "PTY size được cli-gate track dynamic — resize cửa sổ giữa session đều cập nhật word-wrap đúng.")
    add_bullet(doc, "delete vs purge: chọn delete khi muốn ngắt access tạm thời; purge chỉ dùng khi chắc chắn cần xoá hẳn (test account, cleanup).")

    out = "/home/phatlc/data/serverGoChi/docs/cli-config-reference.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
