"""Generate admin web (frontend) reference document (.docx).

Run: python3 docs/build_web_doc.py
Output: docs/web-admin-reference.docx
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ─────────────────────────────────────────────────────────────

def set_style(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)


def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(sh)
    return p


def add_h(doc, text, level): return doc.add_heading(text, level=level)
def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold
    return p
def add_bullet(doc, text): return doc.add_paragraph(text, style="List Bullet")
def add_num(doc, text): return doc.add_paragraph(text, style="List Number")


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            t.rows[ri].cells[ci].text = v
    return t


# ── document ─────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_style(doc)

    title = doc.add_heading("Admin Web — Hướng dẫn sử dụng", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Tham chiếu giao diện admin embedded tại /admin của cli-mgt-svc")
    s.italic = True
    s.font.size = Pt(11)

    # 1. Tổng quan
    add_h(doc, "1. Tổng quan", 1)
    add_para(doc,
             "Admin Web là giao diện quản trị duy nhất cho cli-mgt-svc, embedded sẵn trong "
             "binary (file `pkg/handler/frontend.html` go:embed). Truy cập tại "
             "`http://<host>:3000/admin`. Phần lớn surface chỉ admin/superadmin thấy "
             "(class CSS `.admin-only`); Normal user chỉ thấy Dashboard, History, Guide.")
    add_table(doc,
              ["Đặc điểm", "Chi tiết"],
              [
                  ["Framework", "Plain HTML + JavaScript vanilla — không React/Vue/jQuery; CSS nội tuyến (~250 dòng)"],
                  ["Auth", "JWT lưu localStorage `mgt_token`; gửi qua header `Authorization`"],
                  ["i18n", "EN / VI; chuyển ngôn ngữ ở dropdown sidebar; lưu `mgt_lang`"],
                  ["Theme", "Dark mode duy nhất, palette gradient, accent màu theo category"],
                  ["Sidebar", "5 category — Identity / Network / RBAC / Security / Audit & Tools"],
                  ["No build step", "Sửa frontend.html → rebuild Docker `cli-mgt` là xong"],
              ])

    # 2. Truy cập + đăng nhập
    add_h(doc, "2. Đăng nhập", 1)
    add_num(doc, "Mở `http://<host>:3000/admin` (mặc định cli-mgt-svc lắng nghe :3000).")
    add_num(doc, "Form login hỏi username + password. Gọi `POST /aa/authenticate` để lấy JWT.")
    add_num(doc, "Token được lưu trong `localStorage`. Refresh trang sẽ tự động gọi `/aa/validate-token` — token còn hiệu lực thì vào thẳng app.")
    add_num(doc, "Logout (góc phải topbar) xoá token + reload form login.")
    add_para(doc,
             "Lockout: nếu effective password policy của user bật max_login_failure, "
             "sai pass N lần liên tiếp sẽ bị khoá lockout_minutes phút. Server trả 403 "
             "`{locked_until, retry_in_seconds}` — frontend hiển thị toast lỗi.")

    # 3. Cấu trúc UI
    add_h(doc, "3. Cấu trúc UI", 1)
    add_table(doc,
              ["Vùng", "Mô tả"],
              [
                  ["Sidebar (cố định trái, 230px)", "Brand mark + 5 category header + nav links + language switch. Active link có gradient + leading dot màu theo category."],
                  ["Topbar", "Tiêu đề section hiện tại + tên user đang login + nút My Profile / Change Password / Logout."],
                  ["Main content", "Section đang active. Mỗi section gồm 1+ card; card có left-border màu theo category."],
                  ["Modal", "Dùng cho edit form, manage modal, double-confirm purge. Đóng bằng nút Cancel hoặc click ngoài overlay."],
                  ["Toast", "Góc phải trên — success (xanh) hoặc error (đỏ). Click toast để dismiss sớm."],
              ])

    # 4. Categories + Tabs
    add_h(doc, "4. Sidebar — 5 category × 13 tab", 1)

    add_h(doc, "4.1. Identity", 2)
    add_table(doc, ["Tab", "Chức năng"], [
        ["Users", "CRUD user. Có 4 action button mỗi row: Edit / Reset password (admin) / Disable (soft) / Purge (hard, double-confirm)"],
        ["Groups", "CRUD group + modal Manage để gán tập user + tập NE cho group"],
    ])

    add_h(doc, "4.2. Network", 2)
    add_table(doc, ["Tab", "Chức năng"], [
        ["Network Elements", "CRUD NE (inline edit). Hỗ trợ ne_profile dropdown."],
        ["NE Mapping", "Gán user → NE. Radio chọn target = NE / Group. Bảng tách Direct (badge xanh dương) và Via group (badge xanh lá)."],
        ["NE Profiles", "CRUD profile (SMF, AMF, UPF, …). Profile dùng để gắn nhãn NE cho command-def."],
    ])

    add_h(doc, "4.3. RBAC / Commands (admin-only)", 2)
    add_table(doc, ["Tab", "Chức năng"], [
        ["Command Catalog",
         "Pool pattern thô (lưu localStorage). Import từ XML hoặc CSV. Filter substring + wildcard `show *`. "
         "Multi-select + 'Select all matching'. Bulk Apply: chọn 1 lần metadata (service / ne_profile / category / "
         "risk_level) → áp cho mọi pattern đã tick → batch POST /aa/command-def/import. Workflow 2 bước này thay "
         "cho việc gõ tay từng command-def khi có vendor manifest."],
        ["Command Defs",
         "CRUD command registry + filter (service/profile/category) + CSV import/export. Header CSV: "
         "`service,ne_profile,pattern,category,risk_level,description`. \"Load Sample\" điền 9 row mẫu. "
         "Dùng cho row-by-row CRUD; với bulk thì dùng Command Catalog."],
        ["Command Groups",
         "CRUD bundle. Mỗi group có button \"Manage Commands\" mở modal — chọn def từ dropdown để add/remove member."],
        ["Group Permissions",
         "Chọn group → list rule (effect = ALLOW/DENY badge xanh/đỏ, grant_type = command-group/category/pattern). "
         "Form thêm rule yêu cầu: effect, service, ne_scope, grant_type, grant_value. Revoke 1 rule theo id."],
    ])

    add_h(doc, "4.4. Security (admin-only)", 2)
    add_table(doc, ["Tab", "Chức năng"], [
        ["Password Policies",
         "CRUD policy. Trường: name, min_length, max_age_days (0=never), history_count (0=no check), "
         "max_login_failure (0=no lockout), lockout_minutes, require_{uppercase,lowercase,digit,special}. "
         "Card thứ ba để assign / unassign policy cho 1 group."],
        ["Mgt Permissions",
         "Per-group `(resource, action)` cho các surface mgt-svc: user/ne/group/command/policy/history × "
         "create/read/update/delete. `*` là wildcard 2 phía."],
    ])

    add_h(doc, "4.5. Audit & Tools", 2)
    add_table(doc, ["Tab", "Chức năng"], [
        ["History",
         "Lịch sử thao tác CLI. 6 cột: Time, User, Command, NE, Result, Scope. "
         "Filter theo user/NE/scope/limit. Dòng nào account=`unknown` (downstream service quên gửi) "
         "được badge amber để admin nhận biết."],
        ["Import",
         "Bulk import legacy entities qua text file format `[users]`, `[nes]`, `[user_roles]`, `[user_nes]`. "
         "Cũng có nút Export để dump snapshot."],
        ["Guide", "Tài liệu inline song ngữ EN / VI. Tự cập nhật theo features mới nhất."],
    ])

    # 5. Workflows
    add_h(doc, "5. Workflow chính (admin)", 1)

    add_h(doc, "5.1. Setup ban đầu — onboarding hệ thống mới", 2)
    add_num(doc, "Đăng nhập với SuperAdmin (seed sẵn lúc init DB).")
    add_num(doc, "Vào tab **NE Profiles** → tạo profile cho từng loại NE: SMF, AMF, UPF, …")
    add_num(doc, "Vào tab **Network Elements** → tạo từng NE và gán profile tương ứng.")
    add_num(doc, "Vào tab **Command Catalog** → import vendor XML/CSV manifest vào pool. Filter + multi-select + Bulk Apply để gán metadata (service/profile/category/risk) cho từng nhóm pattern. Tab **Command Defs** dùng cho row-by-row CRUD nếu cần chỉnh tay.")
    add_num(doc, "Vào tab **Command Groups** → gom command-def thành bundle theo profile (vd: smf-subscriber-ops, smf-session-ops).")
    add_num(doc, "Vào tab **Password Policies** → tạo các policy (strict / standard / relaxed).")
    add_num(doc, "Vào tab **Groups** → tạo group cho từng team (team-smf, team-amf, noc-l1, noc-l2, …).")
    add_num(doc, "Cho mỗi group: vào **Group Permissions** → thêm allow/deny rule; vào **Mgt Permissions** → cấp quyền quản trị nếu cần; "
                 "vào tab Password Policies → assign policy.")
    add_num(doc, "Vào tab **Network Elements** hoặc **NE Mapping** → gán NE cho group.")
    add_num(doc, "Vào tab **Users** → tạo user và gán vào group(s) ngay khi tạo (chọn 1 hoặc nhiều group trong form).")

    add_h(doc, "5.2. Onboarding 1 user mới (đã có sẵn group + policy + NE)", 2)
    add_num(doc, "Tab **Users** → bấm Create. Điền username, password, email, full_name, phone, account_type (1=Admin / 2=Normal), chọn group(s).")
    add_num(doc, "User sẽ kế thừa NE access từ group(s), allow/deny rule từ group(s), password policy strict-est merge.")
    add_num(doc, "Báo cho user thông tin login + URL admin (nếu là admin) hoặc URL SSH bastion (nếu là normal).")

    add_h(doc, "5.3. Offboarding 1 user — disable hay purge?", 2)
    add_table(doc, ["Tình huống", "Hành động"], [
        ["User nghỉ tạm, có thể quay lại", "**Disable** (button Disable) — set is_enable=false, dữ liệu giữ nguyên"],
        ["User chỉ tạm cấm vì sự cố", "Disable, mở lại bằng cách tạo lại user cùng username → re-enable + merge field mới"],
        ["Test account, sandbox, tài khoản tạm", "**Purge** — xoá hẳn"],
        ["Credential bị compromise, audit yêu cầu xoá hết dấu vết", "**Purge**. cli_operation_history vẫn được giữ để audit trail còn"],
        ["GDPR / yêu cầu xoá data cá nhân", "**Purge** + cân nhắc xoá luôn các history liên quan qua API"],
    ])
    add_para(doc, "Purge bắt buộc admin gõ lại đúng username để xác nhận — không thể bấm nhầm. SuperAdmin không purge được qua web, phải sửa trực tiếp DB.")

    add_h(doc, "5.4. Cấp quyền cho 1 group thực hiện 1 nhóm lệnh trên 1 loại NE", 2)
    add_num(doc, "Tab **Group Permissions** → chọn group ở dropdown.")
    add_num(doc, "Form Add Permission: effect=allow, service=ne-command, "
                 "ne_scope=`profile:SMF` (áp cho tất cả NE thuộc profile SMF), "
                 "grant_type=command-group, grant_value=`smf-subscriber-ops`.")
    add_num(doc, "Bấm Add → rule xuất hiện ở bảng Current Permissions với badge xanh ALLOW.")
    add_num(doc, "Để cấm thêm 1 NE production cụ thể: thêm rule deny với `ne_scope=ne:HTSMF-PROD-01`. "
                 "Scope ne:X cụ thể hơn profile:SMF nên override cho NE đó.")

    add_h(doc, "5.5. Đổi mật khẩu (user tự đổi)", 2)
    add_num(doc, "Topbar → **Change Password**.")
    add_num(doc, "Modal hỏi old + new + confirm. Server validate theo effective policy: min_length, require_*, history_count.")
    add_num(doc, "Nếu fail policy, toast hiện thông báo cụ thể (vd: 'password must contain a digit').")
    add_num(doc, "Đổi xong, password_expires_at được set lại theo `now + max_age_days`. login_failure_count reset về 0.")

    add_h(doc, "5.6. Admin reset pass cho user khác", 2)
    add_num(doc, "Tab **Users** → row của target → Reset Password.")
    add_num(doc, "Modal yêu cầu nhập new password. Bỏ qua check old-password nhưng vẫn validate theo policy + history.")

    # 6. Tài nguyên đặc biệt
    add_h(doc, "6. Tài nguyên đặc biệt", 1)

    add_h(doc, "6.0. Command Catalog (workflow 2-bước)", 2)
    add_para(doc, "Tab Command Catalog có 3 card xếp dọc:")
    add_bullet(doc, "**1. Import patterns** — upload file (.xml/.csv) hoặc paste vào textarea. Auto-detect format theo ký tự đầu (`<` = XML). XML parser quét bất kỳ element <command>/<cmd>/<pattern> với text hoặc attribute pattern/name/value. CSV parser cần header `pattern[,description]`. Pool ghi vào localStorage `mgt_command_catalog`.")
    add_bullet(doc, "**2. Pool** — table tất cả pattern hiện có. Filter input chấp nhận substring (`get`) hoặc wildcard suffix (`show *`). \"Select all matching\" tick toàn bộ row hiển thị; checkbox header tick/untick visible. Mỗi row có nút × để xoá khỏi pool.")
    add_bullet(doc, "**3. Bulk Apply** — chọn service / ne_profile / category / risk_level 1 lần → bấm Apply → frontend build array {service, ne_profile, pattern, category, risk_level, description} cho mọi pattern đã tick → POST /aa/command-def/import. Backend dedupe theo (service, ne_profile, pattern), nên import lại không tạo duplicate.")
    add_para(doc, "Workflow: import 1 lần lúc onboard vendor → mỗi khi cần phân loại 1 nhóm command thì filter + select + bulk apply, không phải gõ tay từng pattern.")
    add_para(doc, "Pool LOCAL theo browser — admin khác không thấy. Để share, download CSV → gửi cho người khác → họ paste lại.")

    add_h(doc, "6.1. CSV import/export Command Defs", 2)
    add_para(doc, "Tab Command Defs có riêng card \"CSV Import / Export\":")
    add_bullet(doc, "**Download CSV** — xuất tất cả command-def đang có (sau khi áp filter service/profile/category) ra file CSV.")
    add_bullet(doc, "**Upload CSV** — chọn file local, parse client-side, POST batch tới `/aa/command-def/import`.")
    add_bullet(doc, "**Load Sample** — điền sẵn 9 row mẫu vào textarea.")
    add_para(doc, "Header CSV bắt buộc: `service,ne_profile,pattern,category,risk_level,description`. Giá trị chứa dấu phẩy phải đặt trong dấu nháy kép, escape `\"\"` cho nháy kép trong giá trị.")
    add_code(doc,
             "service,ne_profile,pattern,category,risk_level,description\n"
             "ne-command,*,show version,monitoring,0,Show device version\n"
             "ne-command,*,ping *,monitoring,0,\n"
             "ne-command,SMF,get subscriber,monitoring,0,\n"
             "ne-command,SMF,\"delete session *\",admin,2,Destructive: remove PDU session\n"
             "ne-command,AMF,show ue *,monitoring,0,\n"
             "ne-config,*,get-config *,monitoring,0,\n"
             "ne-config,*,edit-config *,configuration,1,")

    add_h(doc, "6.2. Manage Commands modal (Command Groups)", 2)
    add_para(doc, "Bấm \"Manage Commands\" trên row của 1 command group:")
    add_bullet(doc, "Phần trên: list current member với button `-` để remove.")
    add_bullet(doc, "Phần dưới: dropdown chọn def chưa thuộc group, bấm Add để add member mới.")
    add_bullet(doc, "Reload modal sau mỗi action để hiện state mới.")

    add_h(doc, "6.3. Bulk import legacy entities (tab Import)", 2)
    add_table(doc, ["Section", "Trường"], [
        ["[users]", "username, password, email (tùy chọn)"],
        ["[nes]", "ne_name, site_name, namespace, command_url, conf_mode, conf_master_ip, conf_port_master_ssh, conf_username, conf_password, description"],
        ["[user_roles]", "username, permission (admin hoặc user)"],
        ["[user_nes]", "username, ne_name"],
    ])
    add_para(doc, "Existing user/role được skip (không trùng). Export dùng password placeholder `changeme` vì password đã bcrypt-hash không recover được.")

    # 7. Phân quyền hiển thị tab
    add_h(doc, "7. Phân quyền hiển thị tab", 1)
    add_table(doc,
              ["Tab", "SuperAdmin / Admin", "Normal user"],
              [
                  ["Dashboard", "✅", "✅"],
                  ["Users", "✅ (full CRUD)", "Chỉ thấy mình qua My Profile + Change Password"],
                  ["Groups", "✅", "❌"],
                  ["Network Elements", "✅ (full CRUD)", "❌ (read-only NE list lấy từ /aa/list/ne)"],
                  ["NE Mapping", "✅", "❌"],
                  ["NE Profiles", "✅", "❌"],
                  ["Command Catalog", "✅ (pool localStorage per-browser)", "❌"],
                  ["Command Defs", "✅", "❌"],
                  ["Command Groups", "✅", "❌"],
                  ["Group Permissions", "✅", "❌"],
                  ["Password Policies", "✅", "❌"],
                  ["Mgt Permissions", "✅", "❌"],
                  ["History", "✅ (xem all)", "✅ (chỉ thao tác của mình — server filter theo account)"],
                  ["Import", "✅", "❌"],
                  ["Guide", "✅", "✅"],
              ])
    add_para(doc, "Cơ chế: HTML element có class `.admin-only` bị ẩn khi `body.role-user` (set khi JWT có aud=user).")

    # 8. UX behaviors
    add_h(doc, "8. Hành vi UX cần biết", 1)
    add_bullet(doc, "Toast tự biến mất sau 3 giây (success) hoặc 6 giây (error). Click toast để dismiss sớm.")
    add_bullet(doc, "Hover row trong table có hiệu ứng background nhẹ. Hover button primary/danger có lift 1px + shadow.")
    add_bullet(doc, "Form input focus có ring xanh 3px.")
    add_bullet(doc, "Modal đóng bằng nút Cancel/Close (không hỗ trợ Esc — phải click).")
    add_bullet(doc, "Autocomplete dropdown (NE Mapping, History filter): mũi tên ↑/↓ + Enter để chọn, Esc để đóng.")
    add_bullet(doc, "Section switch có animation fadeIn 0.2s.")
    add_bullet(doc, "Resize cửa sổ trình duyệt: layout responsive minimum, không hỗ trợ mobile thực sự.")

    # 9. Troubleshooting
    add_h(doc, "9. Troubleshooting", 1)
    add_table(doc,
              ["Triệu chứng", "Nguyên nhân", "Cách xử lý"],
              [
                  ["Login → 401", "Sai password hoặc account disabled", "Kiểm tra is_enable + admin reset pass nếu cần"],
                  ["Login → 403 \"locked\"", "Vượt max_login_failure", "Chờ hết lockout_minutes hoặc admin reset password (action reset đã reset luôn counter)"],
                  ["Tab nào đó hiển thị empty", "User không có quyền đọc data hoặc DB chưa migrate", "Check role + chạy lại auto-migrate (nếu MySQL/Postgres mới)"],
                  ["Frontend Guide trống", "i18n chưa load", "Reload trang; check console JS"],
                  ["Tab RBAC không thấy", "Login với Normal user", "Cần Admin/SuperAdmin"],
                  ["CSV import báo lỗi parse", "Header sai hoặc field quote chưa đúng", "Check sample, escape `\"\"` cho dấu nháy kép trong value"],
                  ["History dòng nào account=unknown", "Downstream service quên gửi `account` trong body", "Sửa downstream gọi /history/save kèm account"],
                  ["Token validate bị reject sau update mgt-svc", "JWT_SECRET_KEY đổi", "Logout + login lại"],
              ])

    # 10. Tham chiếu API
    add_h(doc, "10. Tham chiếu nhanh API mỗi tab gọi", 1)
    add_table(doc, ["Tab", "Endpoints chính"], [
        ["Dashboard", "GET /aa/admin/user/list, GET /aa/admin/ne/list"],
        ["Users", "POST /aa/authenticate/user/{set,delete,purge,reset-password}, GET /aa/admin/user/full"],
        ["Groups", "GET/POST /aa/group/{list,create,update,delete,show}, /user/assign|unassign, /ne/assign|unassign"],
        ["Network Elements", "GET /aa/admin/ne/list, POST /aa/admin/ne/{create,update}"],
        ["NE Mapping", "POST /aa/authorize/ne/set|delete, /aa/group/user/assign|unassign"],
        ["NE Profiles", "GET /aa/ne-profile/list, POST /aa/ne-profile/{create,update}, DELETE /{id}, POST /aa/ne/{ne_id}/profile"],
        ["Command Defs", "GET /aa/command-def/list, POST /aa/command-def/{create,update,import}, DELETE /{id}"],
        ["Command Groups", "GET /aa/command-group/list, POST /aa/command-group/{create,update}, DELETE /{id}, /aa/command-group/{id}/commands"],
        ["Group Permissions", "GET/POST /aa/group/{id}/cmd-permissions, DELETE /aa/group/{id}/cmd-permissions/{perm_id}"],
        ["Password Policies", "GET /aa/password-policy/list, POST /aa/password-policy/{create,update}, DELETE /{id}, POST /aa/group/{id}/password-policy"],
        ["Mgt Permissions", "GET/POST /aa/group/{id}/mgt-permissions, DELETE /aa/group/{id}/mgt-permissions/{perm_id}"],
        ["History", "GET /aa/history/list (filter via query)"],
        ["Import", "POST /aa/import/"],
    ])

    out = "/home/phatlc/data/serverGoChi/docs/web-admin-reference.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
